from typing import Annotated
from fastapi import FastAPI, File, Form, UploadFile, HTTPException
from youtube_transcript_api import YouTubeTranscriptApi
from python_pptx import Presentation
from docx import Document
from io import BytesIO
from urllib.parse import urlparse, parse_qs


app = FastAPI()


@app.get("/health")
async def health() -> dict[str, str]:
    return {"message": "Healthy!"}

@app.post("/upload")
async def upload(
    files: Annotated[list[UploadFile] | None, File()] = None,
    url: Annotated[str | None, Form()] = None,
):

    files_given = files is not None and len(files) > 0
    url_given = url is not None and url.strip() != ""

    if url_given:
        # Check it's from youtube
        if not url.startswith("https://www.youtube.com/watch?v="):
            raise HTTPException(
                status_code=422,
                detail="Invalid YouTube URL.",
            )

        text = extract_youtube_text(url)
        #TODO: Store in DB

    if not files_given and not url_given:
        raise HTTPException(
            status_code=422,
            detail="Provide at least one file or a YouTube URL.",
        )

    if files_given:
        for f in files:
            if not f.filename.lower().endswith((".docx", ".pptx")):
                raise HTTPException(
                    status_code=415,
                    detail=f"Unsupported file type: {f.filename}",
                )
            else:
                if f.filename.lower().endswith(".docx"):
                    text = extract_docx_text(f.file)
                elif f.filename.lower().endswith(".pptx"):
                    text = extract_pptx_text(f.file)
            

    return {"message": "OK"}

def extract_youtube_text(url: str):
    parsed_url = urlparse(url)
    video_id = parse_qs(parsed_url.query).get('v', [None])[0]
    ytt_api = YouTubeTranscriptApi()
    fetched_transcript = ytt_api.fetch_transcript(video_id)
    text = []
    for t in fetched_transcript:
        text.append(t["text"])

    return " ".join(text)


def extract_pptx_text(file_path) :
    prs = Presentation(BytesIO(file_path.read()))

    text_runs = []

    for slide in prs.slides:
        for shape in slide.shapes:
            if not shape.has_text_frame:
                continue
            for paragraph in shape.text_frame.paragraphs:
                for run in paragraph.runs:
                    text_runs.append(run.text)
    
    return " ".join(text_runs)

def extract_docx_text(file_path) -> str:
    doc = Document(BytesIO(file_path.read()))
    text = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text.append(cell.text)

    return " ".join(text)